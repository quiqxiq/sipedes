# -*- coding: utf-8 -*-
"""
Build BAB IV & BAB V (skripsi SI Pelayanan Desa Rombiyah Barat) as .docx
matching BAB-IV-V-SIPADES.docx formatting exactly: A4, margins 4/3/3/3 cm,
1.5 line spacing (line=360), TNR, Heading1 (BAB) centered bold, Heading2/3
numbered, justified body with firstLine indent, centered figure captions,
shaded table headers, PAGE footer.
"""
import os, sys
from PIL import Image
from docx import Document
from docx.shared import Twips, Pt, Inches, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.abspath(os.path.join(BASE, "..", "screenshots"))
OUT = os.path.abspath(os.path.join(BASE, "..", "screenshots", "BAB-IV-V-SIPEDES-RombiyahBarat-DaftarGambar.docx"))

# Daftar seluruh gambar (urutan harus sama dengan pemanggilan add_figure di generate_docx)
FIGURES = [
    ("01_publik_landing.png", "Halaman Utama / Landing Page Publik SIPEDES Rombiyah Barat"),
    ("02_login_warga.png", "Antarmuka Halaman Login Portal Warga"),
    ("03_registrasi_warga.png", "Formulir Registrasi Mandiri Warga Desa Rombiyah Barat"),
    ("05_warga_dashboard.png", "Halaman Dashboard Utama Portal Warga"),
    ("06_warga_pengajuan_surat.png", "Wizard Pengajuan Surat - Langkah 1 Pemilihan Jenis Surat"),
    ("07_warga_riwayat_index.png", "Halaman Riwayat Pengajuan Surat Portal Warga"),
    ("08_warga_riwayat_detail.png", "Halaman Detail Pengajuan & Status Permohonan Warga"),
    ("04_admin_login.png", "Halaman Login Panel Admin & Perangkat Desa"),
    ("10_admin_dashboard.png", "Dashboard Utama Panel Admin Filament v5"),
    ("11_admin_jenis_surat_index.png", "Halaman Daftar Jenis Surat Pelayanan Desa"),
    ("12_admin_jenis_surat_create.png", "Formulir Tambah Jenis Surat & Persyaratan Berkas"),
    ("13_admin_jenis_surat_edit.png", "Formulir Edit Jenis Surat"),
    ("14_admin_knowledge_documents_index.png", "Halaman Daftar Dokumen Basis Pengetahuan (Dify RAG)"),
    ("15_admin_knowledge_documents_create.png", "Formulir Unggah Dokumen Basis Pengetahuan"),
    ("16_admin_knowledge_documents_edit.png", "Formulir Edit Dokumen Basis Pengetahuan"),
    ("17_admin_permohonan_surat_index.png", "Daftar Permohonan Surat Masuk pada Panel Admin"),
    ("18_admin_permohonan_surat_view.png", "Halaman Detail & Verifikasi Permohonan Surat"),
    ("19_admin_permohonan_surat_edit.png", "Formulir Edit Data Permohonan Surat"),
    ("20_admin_profil_desa_index.png", "Halaman Data Profil Desa Rombiyah Barat"),
    ("21_admin_profil_desa_edit.png", "Formulir Edit Profil Desa"),
    ("22_admin_users_index.png", "Halaman Manajemen Pengguna Sistem"),
    ("23_admin_users_create.png", "Formulir Tambah Pengguna Baru"),
    ("24_admin_users_edit.png", "Formulir Edit Data Pengguna"),
    ("25_admin_aktivitas_logs_index.png", "Halaman Log Aktivitas Sistem (Audit Trail)"),
]

# Lebar area teks dalam twips: 11907 - 2268 (kiri) - 1701 (kanan) = 7938
TEXT_WIDTH_TWIPS = 7938

# ---------------- low-level formatting helpers ----------------

def _child(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def set_spacing(p, before=None, after=0, line=360, rule="auto"):
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), rule)

def set_indent(p, firstLine=None, left=None, hanging=None):
    pPr = p._p.get_or_add_pPr()
    ind = _child(pPr, "w:ind")
    if firstLine is not None:
        ind.set(qn("w:firstLine"), str(firstLine))
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if hanging is not None:
        ind.set(qn("w:hanging"), str(hanging))

def style_run(r, pt=12, bold=False, italic=False, color=None, name="Times New Roman"):
    r.font.name = name
    rpr = r._r.get_or_add_rPr()
    rf = _child(rpr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)
    r.font.size = Pt(pt)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)

# ---------------- page setup ----------------

def setup_page(doc):
    """A4, margins: top/right/bottom=3cm, left=4cm; footer distance 1.27cm."""
    sec = doc.sections[0]
    sec.page_width  = Twips(11907)
    sec.page_height = Twips(16840)
    sec.top_margin    = Twips(1701)
    sec.right_margin  = Twips(1701)
    sec.bottom_margin = Twips(1701)
    sec.left_margin   = Twips(2268)
    sec.footer_distance = Twips(720)
    sec.header_distance = Twips(720)

# ---------------- footer (centered PAGE field) ----------------

def add_page_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    sdtId = OxmlElement("w:id"); sdtId.set(qn("w:val"), "-264543138")
    sdtPr.append(sdtId)
    dpo = OxmlElement("w:docPartObj")
    dpg = OxmlElement("w:docPartGallery"); dpg.set(qn("w:val"), "Page Numbers (Bottom of Page)")
    dpu = OxmlElement("w:docPartUnique")
    dpo.append(dpg); dpo.append(dpu); sdtPr.append(dpo)
    sdt.append(sdtPr)
    sdtEnd = OxmlElement("w:sdtEndPr")
    rpr_end = OxmlElement("w:rPr"); np_end = OxmlElement("w:noProof"); rpr_end.append(np_end)
    sdtEnd.append(rpr_end); sdt.append(sdtEnd)
    sdtContent = OxmlElement("w:sdtContent")

    p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "Footer"); pPr.append(pStyle)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    p_el.append(pPr)
    def fld_run(ftype):
        r = OxmlElement("w:r"); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), ftype); r.append(fc); return r
    p_el.append(fld_run("begin"))
    r_instr = OxmlElement("w:r"); it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = " PAGE   \\* MERGEFORMAT "; r_instr.append(it); p_el.append(r_instr)
    p_el.append(fld_run("separate"))
    r_val = OxmlElement("w:r"); rpr_v = OxmlElement("w:rPr"); np_v = OxmlElement("w:noProof")
    rpr_v.append(np_v); r_val.append(rpr_v)
    t_v = OxmlElement("w:t"); t_v.text = "1"; r_val.append(t_v); p_el.append(r_val)
    p_el.append(fld_run("end"))
    sdtContent.append(p_el); sdt.append(sdtContent)
    footer._element.append(sdt)

# ---------------- paragraph factories ----------------

def add_bab_heading(doc, roman, title):
    """Heading1: 'BAB IV\nHASIL DAN PEMBAHASAN' centered, TNR 12pt bold."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r1 = p.add_run("BAB " + roman); style_run(r1, pt=12, bold=True)
    r_br = p.add_run(); r_br._r.append(OxmlElement("w:br"))
    r2 = p.add_run(title); style_run(r2, pt=12, bold=True)
    return p

def add_h2(doc, num, title):
    """Heading2: '4.1  Hasil Penelitian' — TNR 12pt bold, before=240, after=60."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "240"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_h3(doc, num, title):
    """Heading3: '4.1.1  Lingkungan Implementasi' — TNR 12pt bold, before=180, after=60."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "180"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_body(doc, text, bold_prefix=None, italic=False):
    """Normal body: justified, firstLine=426 (0.75 cm), TNR 12pt, line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_spacing(p, before=0, after=0, line=360)
    set_indent(p, firstLine=426)
    p.alignment = AL.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        style_run(r_pre, pt=12, bold=True, italic=italic)
    r = p.add_run(text)
    style_run(r, pt=12, italic=italic)
    return p

def add_caption(doc, text):
    """Caption: centered, TNR 11pt bold label, before=120 after=120 line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "120"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    ind = _child(pPr, "w:ind"); ind.set(qn("w:firstLine"), "0")

    # Split caption for bold label if starts with Gambar or Tabel
    if text.startswith("Gambar ") or text.startswith("Tabel "):
        parts = text.split("  ", 1)
        if len(parts) == 2:
            r1 = p.add_run(parts[0] + "  "); style_run(r1, pt=11, bold=True, color="000000")
            r2 = p.add_run(parts[1]); style_run(r2, pt=11, bold=False, color="000000")
        else:
            r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    else:
        r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    return p

# ---------------- image insertion ----------------

IMG_COUNTER = {"n": 0}

def add_figure(doc, filename, caption_text):
    IMG_COUNTER["n"] += 1
    idx = IMG_COUNTER["n"] - 1
    assert idx < len(FIGURES) and FIGURES[idx][0] == filename, \
        f"Urutan gambar tidak sinkron dengan FIGURES: {filename}"
    num_str = f"Gambar 4.{IMG_COUNTER['n']}"
    full_caption = f"{num_str}  {caption_text}"
    path = os.path.join(IMG_DIR, filename)
    try:
        with Image.open(path) as im:
            w_px, h_px = im.size
    except Exception:
        w_px, h_px = 1280, 720

    # Max width: 14 cm (~5.51 in)
    max_w_emu = int(5.51 * 914400)
    max_h_emu = int(3.5 * 914400)
    ratio = w_px / h_px
    w_emu = max_w_emu
    h_emu = int(w_emu / ratio)
    if h_emu > max_h_emu:
        h_emu = max_h_emu
        w_emu = int(h_emu * ratio)

    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60, line=360)
    p.alignment = AL.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(w_emu), height=Emu(h_emu))
    add_caption(doc, full_caption)
    return p

# ---------------- DAFTAR GAMBAR (list of figures) ----------------

def add_page_break(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0, line=360)
    r = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    r._r.append(br)
    return p

def add_front_heading(doc, text):
    """Judul bagian awal (DAFTAR GAMBAR): centered, bold, TNR 12pt."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "240")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r = p.add_run(text); style_run(r, pt=12, bold=True)
    return p

def add_daftar_entry(doc, label, caption, page):
    """Satu baris daftar gambar: label + judul ..... halaman (dot leader)."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_spacing(p, before=0, after=0, line=360)
    p.paragraph_format.tab_stops.add_tab_stop(Twips(TEXT_WIDTH_TWIPS), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(f"{label}  {caption}"); style_run(r1, pt=12)
    r2 = p.add_run("\t" + str(page)); style_run(r2, pt=12)
    return p

def add_daftar_gambar(doc, fig_pages=None):
    """Buat halaman DAFTAR GAMBAR di awal dokumen. fig_pages: dict nomor->halaman."""
    add_page_break(doc)
    add_front_heading(doc, "DAFTAR GAMBAR")
    sp = doc.add_paragraph()
    set_spacing(sp, before=0, after=0, line=360)

    # Baris header kolom halaman
    hp = doc.add_paragraph()
    hp.style = doc.styles["Normal"]
    set_spacing(hp, before=0, after=0, line=360)
    hp.paragraph_format.tab_stops.add_tab_stop(Twips(TEXT_WIDTH_TWIPS), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    rh = hp.add_run("\tHalaman"); style_run(rh, pt=12)

    for i, (_fname, caption) in enumerate(FIGURES, 1):
        page = str(fig_pages.get(i, "")) if fig_pages else ""
        add_daftar_entry(doc, f"Gambar 4.{i}", caption, page)

    add_page_break(doc)

# ---------------- black-box test table ----------------

def add_bb_table(doc, tbl_num, caption_text, rows):
    """rows = list of (no, skenario, input, expected, result, status)"""
    add_caption(doc, f"Tabel {tbl_num}  {caption_text}")
    tbl = doc.add_table(rows=1 + len(rows), cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths (Total ~ 14 cm)
    # No: 0.8 cm, Skenario: 2.8 cm, Input: 2.8 cm, Expected: 3.3 cm, Result: 3.3 cm, Status: 1.0 cm
    widths = [Twips(450), Twips(1580), Twips(1580), Twips(1860), Twips(1860), Twips(600)]
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    headers = ["No.", "Skenario Uji", "Data Masukan", "Hasil yang Diharapkan", "Hasil Pengujian", "Status"]
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_spacing(p, before=60, after=60, line=240)
        p.alignment = AL.CENTER
        r = p.add_run(h); style_run(r, pt=10, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_spacing(p, before=60, after=60, line=240)
            p.alignment = AL.CENTER if ci in (0, 5) else AL.JUSTIFY
            r = p.add_run(str(val))
            is_bold = (ci == 5)
            style_run(r, pt=9.5, bold=is_bold)
    return tbl

# ---------------- DOCUMENT GENERATOR MAIN ----------------

def generate_docx(fig_pages=None):
    doc = Document()
    setup_page(doc)
    add_page_footer(doc)

    print("Generating DAFTAR GAMBAR...")
    add_daftar_gambar(doc, fig_pages)

    print("Generating BAB IV...")

    # ---------------- BAB IV ----------------
    add_bab_heading(doc, "IV", "HASIL DAN PEMBAHASAN")

    # 4.1 Hasil Penelitian
    add_h2(doc, "4.1", "Hasil Penelitian")
    add_body(doc, "Hasil penelitian ini adalah terbangunnya Sistem Informasi Pelayanan Desa (SIPEDES) Rombiyah Barat berbasis web yang dirancang untuk mengotomatisasi seluruh alur permohonan surat-menyurat di Kantor Desa Rombiyah Barat, Kecamatan Gandusari, Kabupaten Blitar. Sistem ini memadukan konsep self-service portal bagi warga desa, wizard pengajuan surat multi-langkah, mekanisme verifikasi dan persetujuan permohonan oleh perangkat desa, otomatisasi pembuatan dokumen surat resmi dalam format PDF, serta integrasi Kecerdasan Buatan (AI) berbasis Retrieval-Augmented Generation (RAG) melalui platform Dify sebagai asisten virtual layanan desa.")

    # 4.1.1 Lingkungan Implementasi
    add_h3(doc, "4.1.1", "Lingkungan Implementasi")
    add_body(doc, "Pengembangan dan implementasi sistem SIPEDES Rombiyah Barat dilaksanakan pada lingkungan perangkat keras (hardware) dan perangkat lunak (software) dengan spesifikasi teknis sebagai berikut:")
    add_body(doc, "1. Spesifikasi Perangkat Keras (Hardware): Sistem dikembangkan pada unit PC/Laptop penguji berbasis prosesor modern multi-core, memori utama (RAM) 16 GB, serta media penyimpanan Solid State Drive (SSD) NVMe 512 GB. Lingkungan peladen (server) pengujian dijalankan secara lokal menggunakan web server bawaan Laravel (php artisan serve) pada port 8000, siap dipindahkan ke infrastruktur Cloud Server (VPS) pada tahap produksi.")
    add_body(doc, "2. Spesifikasi Perangkat Lunak (Software): Sistem beroperasi pada lingkungan Sistem Operasi Microsoft Windows 11 Pro. Bahasa pemrograman utama yang digunakan adalah PHP 8.3 dengan basis web framework Laravel 13, panel manajemen administrasi Filament v5, basis data SQLite (pengembangan) dan MySQL 8.0 (produksi), web server Nginx/Apache, serta browser Google Chrome untuk pengujian fungsionalitas.")
    add_body(doc, "3. Tech Stack dan Pustaka Pendukung: Arsitektur perangkat lunak memanfaatkan Blade + Livewire untuk komponen UI dinamis portal warga, Filament v5 untuk panel administrasi dan manajemen resource, Barryvdh Laravel DomPDF untuk pembentukan dokumen surat resmi dalam format PDF, Tailwind CSS 4 dan Vite 8 untuk styling serta bundling aset, serta Dify (open-source LLMOps platform) sebagai layanan Knowledge Base dan RAG chatbot dengan HTTP API streaming.")

    # 4.1.2 Implementasi Antarmuka Portal Warga
    add_h3(doc, "4.1.2", "Implementasi Antarmuka Portal Warga")
    add_body(doc, "Antarmuka portal warga dirancang berorientasi pada kemudahan pengguna (user-friendly) dan kemudahan akses mandiri (self-service). Warga dapat mengakses halaman utama publik untuk melihat informasi umum layanan desa sebelum melakukan autentikasi.")

    add_figure(doc, "01_publik_landing.png", "Halaman Utama / Landing Page Publik SIPEDES Rombiyah Barat")
    add_body(doc, "Halaman utama (Landing Page) menampilkan hero banner selamat datang, daftar informasi layanan persuratan unggulan (Surat Keterangan Tidak Mampu, Domisili, Usaha, Pengantar Nikah, dan Keterangan Kematian), statistik pelayanan, serta tautan menuju halaman login dan registrasi portal warga.")

    add_figure(doc, "02_login_warga.png", "Antarmuka Halaman Login Portal Warga")
    add_body(doc, "Halaman login portal warga digunakan oleh warga terdaftar untuk masuk ke dalam sistem menggunakan Nomor Induk Kependudukan (NIK) dan kata sandi yang telah dibuat saat registrasi mandiri.")

    add_figure(doc, "03_registrasi_warga.png", "Formulir Registrasi Mandiri Warga Desa Rombiyah Barat")
    add_body(doc, "Warga baru yang belum memiliki akun dapat mengisikan formulir registrasi mandiri yang mencakup Nama Lengkap sesuai KTP, NIK (16 digit), Alamat Email, Nomor Telepon/WhatsApp aktif, Alamat Tempat Tinggal, serta Kata Sandi dengan konfirmasi. Sistem memvalidasi panjang NIK dan keunikan data sebelum akun dibuat.")

    add_figure(doc, "05_warga_dashboard.png", "Halaman Dashboard Utama Portal Warga")
    add_body(doc, "Setelah berhasil login, warga disambut pada Halaman Dashboard Warga yang memuat kartu informasi ringkasan pengajuan (Total Permohonan, Sedang Diproses, Surat Disetujui), tombol cepat ajukan surat, serta identitas NIK warga yang sedang login.")

    add_figure(doc, "06_warga_pengajuan_surat.png", "Wizard Pengajuan Surat - Langkah 1 Pemilihan Jenis Surat")
    add_body(doc, "Warga dapat mengajukan permohonan surat melalui wizard tiga langkah. Pada langkah pertama, warga memilih jenis surat yang dibutuhkan beserta estimasi waktu proses, deskripsi layanan, dan syarat berkas yang wajib dipersiapkan.")

    add_figure(doc, "07_warga_riwayat_index.png", "Halaman Riwayat Pengajuan Surat Portal Warga")
    add_body(doc, "Halaman riwayat pengajuan menampilkan seluruh berkas permohonan yang pernah dibuat oleh warga lengkap dengan nomor permohonan, jenis surat, tanggal pengajuan, dan indikator badge status permohonan secara real-time.")

    add_figure(doc, "08_warga_riwayat_detail.png", "Halaman Detail Pengajuan & Status Permohonan Warga")
    add_body(doc, "Warga dapat melihat rincian permohonan beserta status terkini (Diajukan, Diproses, Disetujui, Butuh Koreksi, atau Ditolak) serta catatan dari petugas desa. Surat yang telah disetujui dapat diunduh dalam format PDF melalui route /surat/{id}/pdf.")

    # 4.1.3 Implementasi Antarmuka Panel Admin
    add_h3(doc, "4.1.3", "Implementasi Antarmuka Panel Admin & Perangkat Desa")
    add_body(doc, "Panel administrasi dikembangkan menggunakan Filament v5 untuk memfasilitasi tugas perangkat desa dan administrator sistem dalam mengelola master data, permohonan surat, basis pengetahuan AI, serta pemantauan aktivitas sistem.")

    add_figure(doc, "04_admin_login.png", "Halaman Login Panel Admin & Perangkat Desa")
    add_body(doc, "Halaman login khusus pengelola desa yang menjamin autentikasi aman menggunakan alamat email dan kata sandi. Hanya pengguna dengan peran admin yang dapat mengakses seluruh modul manajemen.")

    add_figure(doc, "10_admin_dashboard.png", "Dashboard Utama Panel Admin Filament v5")
    add_body(doc, "Dashboard Admin menampilkan widget statistik permohonan surat beserta grafik distribusi status permohonan (chart) sehingga perangkat desa dapat memantau antrian pelayanan secara ringkas dan real-time.")

    add_figure(doc, "11_admin_jenis_surat_index.png", "Halaman Daftar Jenis Surat Pelayanan Desa")
    add_body(doc, "Admin dapat mengelola seluruh katalog jenis surat yang dilayani oleh Desa Rombiyah Barat, termasuk pengaturan kode surat (SKTM, SKD, SKU, SKN, SKK), estimasi durasi proses, syarat berkas, serta status aktif layanan.")

    add_figure(doc, "12_admin_jenis_surat_create.png", "Formulir Tambah Jenis Surat & Persyaratan Berkas")
    add_body(doc, "Sistem memungkinkan Admin menambahkan jenis surat baru lengkap dengan deskripsi, estimasi waktu, daftar syarat berkas, dan status aktif sesuai kebutuhan peraturan desa yang berlaku.")

    add_figure(doc, "13_admin_jenis_surat_edit.png", "Formulir Edit Jenis Surat")
    add_body(doc, "Admin dapat memperbarui rincian jenis surat, menyesuaikan syarat dokumen, mengubah estimasi waktu proses, serta mengaktifkan/menonaktifkan layanan surat yang bersangkutan.")

    add_figure(doc, "14_admin_knowledge_documents_index.png", "Halaman Daftar Dokumen Basis Pengetahuan (Dify RAG)")
    add_body(doc, "Modul Dokumen Pengetahuan menampilkan seluruh dokumen yang diunggah sebagai basis pengetahuan AI, lengkap dengan status indexing di Dify Knowledge Base (Diproses/Terindeks/Gagal), jumlah chunk, serta aksi Index ke Dify dan Cek Status.")

    add_figure(doc, "15_admin_knowledge_documents_create.png", "Formulir Unggah Dokumen Basis Pengetahuan")
    add_body(doc, "Admin dapat mengunggah dokumen pedoman/SOP pelayanan desa (format PDF, DOCX, TXT) yang secara otomatis dikirim ke Dify untuk diproses chunking dan indexing sebagai sumber jawaban chatbot AI.")

    add_figure(doc, "16_admin_knowledge_documents_edit.png", "Formulir Edit Dokumen Basis Pengetahuan")
    add_body(doc, "Halaman pembaruan dokumen pengetahuan menampilkan informasi file, status indexing di Dify, jumlah chunks, dan identitas pengunggah. Ketika file diganti, sistem otomatis meng-upload ulang dokumen ke knowledge base Dify.")

    add_figure(doc, "17_admin_permohonan_surat_index.png", "Daftar Permohonan Surat Masuk pada Panel Admin")
    add_body(doc, "Halaman kelola permohonan surat menyajikan tabel antrian permohonan masuk lengkap dengan nomor permohonan, pemohon, jenis surat, dan badge status, serta tab penyaring berdasarkan status (Semua, Diajukan, Diproses, Disetujui, Ditolak, Butuh Koreksi).")

    add_figure(doc, "18_admin_permohonan_surat_view.png", "Halaman Detail & Verifikasi Permohonan Surat")
    add_body(doc, "Perangkat desa dapat memeriksa rincian permohonan, data pemohon, jenis surat, berkas lampiran, dan catatan pemohon sebelum melakukan tindakan verifikasi berupa menyetujui dan memproses, meminta koreksi, atau menolak permohonan disertai alasan resmi.")

    add_figure(doc, "19_admin_permohonan_surat_edit.png", "Formulir Edit Data Permohonan Surat")
    add_body(doc, "Halaman pengeditan data permohonan memungkinkan perangkat desa memperbaiki catatan petugas, melampirkan file PDF surat resmi hasil generate, serta mengubah status permohonan sesuai hasil verifikasi.")

    add_figure(doc, "20_admin_profil_desa_index.png", "Halaman Data Profil Desa Rombiyah Barat")
    add_body(doc, "Modul Profil Desa menampilkan informasi identitas wilayah (nama desa, kecamatan, kabupaten, provinsi), sejarah desa, visi & misi, kontak resmi, serta jam operasional layanan kantor desa.")

    add_figure(doc, "21_admin_profil_desa_edit.png", "Formulir Edit Profil Desa")
    add_body(doc, "Admin dapat memperbarui seluruh informasi profil desa, kontak resmi (telepon, WhatsApp, email, alamat kantor), serta jam operasional layanan yang ditampilkan pada portal warga.")

    add_figure(doc, "22_admin_users_index.png", "Halaman Manajemen Pengguna Sistem")
    add_body(doc, "Kelola data akun pengguna sistem yang mencakup Administrator Desa dan seluruh warga terdaftar, lengkap dengan NIK, nomor telepon, status akun aktif, serta peran (role) masing-masing pengguna.")

    add_figure(doc, "23_admin_users_create.png", "Formulir Tambah Pengguna Baru")
    add_body(doc, "Formulir pendaftaran pengguna baru (warga maupun admin) dengan pengisian NIK, nama lengkap, alamat, kontak, kata sandi, serta penetapan peran dan status aktif akun.")

    add_figure(doc, "24_admin_users_edit.png", "Formulir Edit Data Pengguna")
    add_body(doc, "Halaman perbaikan data pengguna untuk memperbarui identitas, kontak, peran, serta mengaktifkan/menonaktifkan akun warga.")

    add_figure(doc, "25_admin_aktivitas_logs_index.png", "Halaman Log Aktivitas Sistem (Audit Trail)")
    add_body(doc, "Pencatatan jejak audit (audit trail) otomatis yang mencatat setiap aksi penting pengguna seperti login, pengajuan surat, persetujuan permohonan, upload dokumen pengetahuan, dan perubahan data, demi menjamin akuntabilitas dan transparansi layanan desa.")

    # 4.1.4 Implementasi Modul Utama & Logika Sistem
    add_h3(doc, "4.1.4", "Implementasi Modul Utama & Logika Sistem")
    add_body(doc, "Sistem SIPEDES Rombiyah Barat mengimplementasikan beberapa modul utama yang saling terintegrasi secara dinamis:")
    add_body(doc, "1. Modul Self-Service Autentikasi Warga: Warga melakukan pendaftaran akun mandiri dengan validasi NIK 16 digit, kemudian masuk ke portal menggunakan kombinasi NIK dan kata sandi. Sesi autentikasi dilindungi oleh mekanisme CSRF dan middleware Laravel, dengan pengalihan otomatis (redirect) bagi pengunjung yang belum login ke halaman login portal warga.")
    add_body(doc, "2. Modul Wizard Pengajuan Surat Multi-Langkah: Pengajuan surat dibangun menggunakan komponen Livewire dengan tiga tahapan, yaitu (1) pemilihan jenis surat, (2) pengunggahan berkas persyaratan beserta catatan keperluan, dan (3) konfirmasi ringkasan data sebelum pengiriman. Validasi berkas (format PDF/JPG/PNG dan ukuran maksimal 3 MB) dilakukan pada sisi server, kemudian sistem membangkitkan nomor permohonan otomatis berformat SRT/YYYYMMDD/XXXXX, menyimpan berkas pada storage publik, dan membuat notifikasi internal untuk warga.")
    add_body(doc, "3. Modul Asisten AI RAG (Dify): Sistem terintegrasi dengan platform Dify melalui App\\Services\\DifyService menggunakan HTTP client Guzzle. Dokumen pedoman dan SOP layanan desa yang diunggah pada modul Dokumen Pengetahuan dikirim ke Knowledge Base Dify untuk diproses (chunking & indexing) dengan pemantauan status secara berkala. Chatbot asisten AI (Livewire ChatbotWidget) tertanam pada setiap halaman portal warga, mampu menerima pertanyaan terkait syarat surat dan SOP pelayanan, menjawab secara streaming real-time dengan mencantumkan sumber dokumen (sources), serta menyimpan riwayat percakapan (ChatSession dan ChatHistory) per warga.")
    add_body(doc, "4. Modul Pembuatan Dokumen Surat PDF: Surat resmi yang telah disetujui dibangkitkan otomatis dalam format PDF menggunakan pustaka Barryvdh Laravel DomPDF melalui controller SuratPdfController, dengan rute pengunduhan /surat/{id}/pdf yang hanya dapat diakses oleh warga pemilik permohonan setelah login.")
    add_body(doc, "5. Modul Manajemen Permohonan & Audit Trail: Perangkat desa memverifikasi permohonan melalui aksi Setujui & Proses, Minta Koreksi, atau Tolak yang masing-masing mencatat alasan ke dalam catatan_petugas dan memperbarui status permohonan. Seluruh tindakan penting dicatat ke dalam tabel aktivitas_log melalui fasilitas AktivitasLog::catat sehingga jejak audit sistem terdokumentasi secara lengkap.")

    # 4.2 Pengujian Sistem
    add_h2(doc, "4.2", "Pengujian Sistem (Testing & Evaluation)")

    # 4.2.1 Metode Pengujian
    add_h3(doc, "4.2.1", "Metode Pengujian")
    add_body(doc, "Pengujian Sistem Informasi Pelayanan Desa Rombiyah Barat (SIPEDES) dilakukan menggunakan metode Black-Box Testing dan Automated End-to-End (E2E) Functional Testing. Pengujian Black-Box berfokus pada evaluasi fungsionalitas antarmuka dan alur kerja sistem tanpa harus melihat struktur kode internal, guna memastikan seluruh kebutuhan fungsional (Functional Requirements) telah terpenuhi dengan benar.")
    add_body(doc, "Pengujian otomatis dilaksanakan menggunakan Puppeteer (Node.js) dengan browser Google Chrome headless terhadap aplikasi yang berjalan pada http://127.0.0.1:8000. Setiap kasus uji membuka halaman, mengisi formulir, dan memverifikasi respons HTTP, konten halaman, pengalihan (redirect), hingga validasi sisi server. Sebelum pengujian, disiapkan data uji berupa 17 warga hasil seeding, 5 jenis surat aktif, 4 contoh permohonan surat, serta 1 dokumen basis pengetahuan sehingga seluruh halaman dapat dirender secara realistis.")

    # 4.2.2 Hasil Pengujian Fungsional
    add_h3(doc, "4.2.2", "Hasil Pengujian Fungsional")
    add_body(doc, "Pengujian fungsionalitas dilakukan mencakup 29 kasus uji utama (Test Cases TC-01 s/d TC-29) yang merepresentasikan seluruh alur proses bisnis halaman publik, portal warga, dan panel admin desa. Hasil pengujian disajikan pada Tabel 4.1 dan Tabel 4.2 berikut.")

    rows_warga = [
        ("TC-01", "Akses Halaman Landing Publik", "GET /", "HTTP 200 dengan judul & konten hero halaman tampil", "Halaman landing tampil lengkap (HTTP 200)", "PASS"),
        ("TC-02", "Akses Halaman Login Warga", "GET /login", "Form NIK & password tersedia", "Form login NIK + password tampil", "PASS"),
        ("TC-03", "Akses Halaman Registrasi Mandiri", "GET /register", "Form 7 field (nama, NIK, email, telepon, alamat, password, konfirmasi)", "Seluruh field registrasi tampil", "PASS"),
        ("TC-04", "Login Warga Password Salah", "NIK valid + password salah", "Muncul pesan error & tetap di halaman login", "Pesan error 'tidak cocok' tampil", "PASS"),
        ("TC-05", "Login Warga Berhasil", "NIK 3529102904650001 + password benar", "Redirect ke /dashboard", "Redirect ke dashboard warga", "PASS"),
        ("TC-06", "Dashboard Warga & Statistik", "GET /dashboard", "Kartu statistik & identitas NIK tampil", "Statistik permohonan & NIK tampil", "PASS"),
        ("TC-07", "Wizard Langkah 1: Pilih Jenis Surat", "Klik kartu SKTM pada /pengajuan", "Pindah ke Langkah 2 (berkas & keterangan)", "Langkah 2 tampil", "PASS"),
        ("TC-08", "Wizard: Kirim Permohonan Lengkap", "Pilih jenis surat, upload PDF, klik Kirim Permohonan", "Redirect ke /riwayat/{id} + notifikasi sukses", "Permohonan tersimpan, redirect & notifikasi sukses", "PASS"),
        ("TC-09", "Halaman Riwayat Pengajuan", "GET /riwayat", "Data permohonan + badge status tampil", "Daftar permohonan & status tampil", "PASS"),
        ("TC-10", "Detail Pengajuan & Status", "GET /riwayat/1", "Detail surat & badge status tampil", "Detail SKTM + badge Disetujui tampil", "PASS"),
        ("TC-11", "Unduh PDF Surat Resmi", "GET /surat/1/pdf", "Respons PDF valid (application/pdf)", "HTTP 200, magic %PDF, content-type pdf", "PASS"),
        ("TC-12", "Logout Warga", "POST /logout", "Sesi berakhir & redirect ke halaman publik", "Redirect ke halaman publik", "PASS"),
    ]
    add_bb_table(doc, "4.1", "Hasil Pengujian Fungsional Halaman Publik & Portal Warga", rows_warga)

    rows_admin = [
        ("TC-13", "Akses Halaman Login Panel Admin", "GET /admin/login", "Form email & password tersedia", "Form login admin tampil", "PASS"),
        ("TC-14", "Login Admin Gagal", "Email benar + password salah", "Muncul pesan error kredensial", "Pesan kredensial tidak ditemukan tampil", "PASS"),
        ("TC-15", "Login Admin Berhasil", "admin@rombiyahbarat.desa.id + password", "Redirect ke dashboard admin", "Redirect ke /admin", "PASS"),
        ("TC-16", "Dashboard Admin Widget & Chart", "GET /admin", "Widget statistik & elemen chart tampil", "Widget + chart SVG tampil", "PASS"),
        ("TC-17", "Jenis Surat: Daftar 5 Jenis Surat", "GET /admin/jenis-surats", "Kode SKTM, SKD, SKU, SKN, SKK tampil", "Kelima kode surat terdeteksi", "PASS"),
        ("TC-18", "Jenis Surat: Formulir Tambah", "GET /admin/jenis-surats/create", "Form input kode & deskripsi tersedia", "Form tampil", "PASS"),
        ("TC-19", "Jenis Surat: Edit Data Terisi", "GET /admin/jenis-surats/1/edit", "Field kode terisi nilai yang tersimpan", "Kode 'SKTM' terisi", "PASS"),
        ("TC-20", "Knowledge Documents: Daftar Dokumen", "GET /admin/knowledge-documents", "Dokumen 'SOP Pelayanan' tampil", "Daftar dokumen pengetahuan tampil", "PASS"),
        ("TC-21", "Knowledge Documents: Formulir Unggah", "GET /admin/knowledge-documents/create", "Form upload file tersedia", "Form unggah tampil", "PASS"),
        ("TC-22", "Permohonan Surat: Daftar Masuk", "GET /admin/permohonan-surats", "4 record contoh permohonan tampil", "4 record permohonan terdeteksi", "PASS"),
        ("TC-23", "Permohonan Surat: Detail/Verifikasi", "GET /admin/permohonan-surats/1", "Nomor permohonan tampil", "SRT/20260813/00001 tampil", "PASS"),
        ("TC-24", "Permohonan Surat: Formulir Edit", "GET /admin/permohonan-surats/1/edit", "Form input/select/textarea tersedia", "Form edit tampil", "PASS"),
        ("TC-25", "Profil Desa: Data Terisi & Edit", "GET /admin/profil-desas + /1/edit", "Nama desa terisi & form edit tersedia", "'Desa Rombiyah Barat' terisi", "PASS"),
        ("TC-26", "Users: Daftar Pengguna", "GET /admin/users", "Admin & warga terdaftar tampil", "Administrator & warga SUHRAWI tampil", "PASS"),
        ("TC-27", "Users: Formulir Tambah Pengguna", "GET /admin/users/create", "Form input nama/role tersedia", "Form tampil", "PASS"),
        ("TC-28", "Aktivitas Logs: Audit Trail", "GET /admin/aktivitas-logs", "Halaman log aktivitas ter-render", "Halaman log render tanpa error", "PASS"),
        ("TC-29", "Registrasi: Validasi Panjang NIK", "NIK 3 digit pada form registrasi", "Muncul pesan error NIK harus 16 digit", "Pesan error validasi NIK tampil", "PASS"),
    ]
    add_bb_table(doc, "4.2", "Hasil Pengujian Fungsional Panel Admin & Validasi Sistem", rows_admin)

    # 4.2.3 Analisis Hasil Pengujian
    add_h3(doc, "4.2.3", "Analisis Hasil Pengujian")
    add_body(doc, "Berdasarkan hasil pengujian fungsionalitas yang dirangkum pada Tabel 4.1 dan Tabel 4.2, dapat ditarik beberapa poin analisis penting:")
    add_body(doc, "1. Tingkat Keberhasilan Fungsional (Success Rate): Dari total 29 skenario pengujian fungsional yang dieksekusi secara otomatis, seluruh kasus uji (100%) dinyatakan PASS (Berhasil) setelah proses perbaikan. Hal ini membuktikan bahwa seluruh fitur utama sistem beroperasi secara stabil dan sesuai dengan spesifikasi kebutuhan.")
    add_body(doc, "2. Penemuan dan Perbaikan Bug Selama Pengujian: Pengujian otomatis berhasil menemukan satu kesalahan nyata (real bug) pada alur pengiriman permohonan wizard, yaitu kegagalan HTTP 500 saat penyimpanan berkas persyaratan karena nama kolom permohonan_surat_id tidak sesuai dengan skema basis data yang menggunakan kolom permohonan_id pada tabel dokumen_persyaratan dan notifikasi. Kesalahan tersebut segera diperbaiki pada komponen PengajuanSuratWizard, kemudian seluruh rangkaian pengujian dijalankan ulang dan dinyatakan lulus (TC-08 PASS). Temuan ini menunjukkan efektivitas metode pengujian black-box otomatis dalam mendeteksi cacat yang tidak terlihat pada tahap pengembangan.")
    add_body(doc, "3. Keandalan Alur Pengajuan & Verifikasi: Pengujian menunjukkan bahwa alur pengajuan tiga langkah (pilih jenis surat, unggah berkas, konfirmasi) berjalan konsisten dengan validasi berkas yang tepat, pembangkitan nomor permohonan otomatis, serta pembaruan status yang tampil secara real-time pada halaman riwayat warga dan daftar permohonan admin.")
    add_body(doc, "4. Integritas Dokumen dan Keamanan Akses: Pembuatan PDF surat resmi menghasilkan dokumen valid (magic bytes %PDF) yang hanya dapat diakses oleh warga pemilik permohonan yang telah login. Halaman admin terlindungi autentikasi Filament dengan pesan error yang informatif, dan seluruh tindakan tercatat pada audit trail aktivitas log.")

    # 4.3 Pembahasan
    add_h2(doc, "4.3", "Pembahasan")

    # 4.3.1 Analisis Efisiensi Pelayanan Publik Desa
    add_h3(doc, "4.3.1", "Analisis Efisiensi Pelayanan Publik Desa")
    add_body(doc, "Implementasi SIPEDES membawa perubahan signifikan terhadap efisiensi dan efektivitas pelayanan administrasi surat-menyurat di Desa Rombiyah Barat. Pada alur konvensional (manual), warga harus datang secara fisik ke Kantor Desa, mengisi formulir kertas, membawa fotokopi berkas persyaratan, dan menunggu antrean pemrosesan oleh perangkat desa. Proses manual ini memakan waktu antara 1 hingga 3 hari kerja tergantung ketersediaan petugas.")
    add_body(doc, "Dengan hadirnya SIPEDES berbasis web self-service, warga dapat mengajukan permohonan dari mana saja dan kapan saja melalui wizard yang membimbing proses pengisian secara bertahap sehingga meminimalkan kesalahan administrasi. Status pengajuan dapat dipantau secara real-time melalui halaman riwayat tanpa perlu melakukan konfirmasi manual ke kantor desa, dan surat yang telah disetujui dapat diunduh langsung dalam format PDF. Otomatisasi penomoran permohonan dan pencatatan audit trail mempercepat kerja perangkat desa dalam memverifikasi berkas, sehingga durasi pelayanan menjadi lebih singkat dan dapat diprediksi.")

    # 4.3.2 Analisis Integrasi AI RAG (Dify) dan Kualitas Informasi Layanan
    add_h3(doc, "4.3.2", "Analisis Integrasi AI RAG (Dify) dan Kualitas Informasi Layanan")
    add_body(doc, "Salah satu keunggulan sistem ini adalah integrasi asisten virtual berbasis AI RAG melalui platform Dify. Dokumen pedoman dan SOP pelayanan yang dikelola oleh admin dijadikan basis pengetahuan (Knowledge Base) yang diproses melalui mekanisme chunking dan indexing, sehingga chatbot mampu memberikan jawaban yang kontekstual sesuai dokumen resmi desa, bukan sekadar jawaban generik. Setiap jawaban dilengkapi dengan sumber (sources) dokumen sehingga warga dapat memverifikasi kebenaran informasi.")
    add_body(doc, "Pengelolaan siklus hidup dokumen pengetahuan (upload otomatis saat pembuatan, re-index saat penggantian file, penghapusan saat dokumen dihapus, serta pemantauan status indexing) menjamin bahwa basis pengetahuan AI selalu sinkron dengan kebijakan layanan desa terkini. Alur percakapan warga dengan asisten AI juga tersimpan dalam riwayat (ChatSession dan ChatHistory) sehingga pengalaman layanan menjadi personal dan berkelanjutan antar sesi.")

    # 4.3.3 Kelebihan dan Keterbatasan Sistem
    add_h3(doc, "4.3.3", "Kelebihan dan Keterbatasan Sistem")
    add_body(doc, "Sistem Informasi Pelayanan Desa (SIPEDES) Rombiyah Barat memiliki beberapa keunggulan utama, antara lain: kemudahan akses mandiri warga melalui portal self-service berbasis NIK, wizard pengajuan tiga langkah yang ramah pengguna, otomatisasi pembuatan dokumen surat resmi PDF, manajemen permohonan dan verifikasi yang terpusat pada panel admin Filament, asisten AI RAG yang menjawab pertanyaan layanan berdasarkan dokumen resmi desa, serta pencatatan audit trail yang menjamin akuntabilitas.")
    add_body(doc, "Meskipun demikian, sistem ini masih memiliki beberapa keterbatasan, di antaranya: operasional sistem sangat bergantung pada ketersediaan koneksi jaringan internet serta ketersediaan layanan Dify sebagai penyedia Knowledge Base AI. Proses indexing dokumen di Dify bersifat asinkron sehingga status indeksasi perlu dipantau secara berkala oleh admin. Selain itu, tanda tangan pada dokumen surat saat ini masih berupa konfirmasi administratif oleh perangkat desa, belum menggunakan Tanda Tangan Elektronik (TTE) bersertifikat dari Penyelenggara Sertifikat Elektronik (PSrE) resmi seperti BSrE BSSN.")

    print("Generating BAB V...")

    # ---------------- BAB V ----------------
    add_bab_heading(doc, "V", "KESIMPULAN DAN SARAN")

    # 5.1 Kesimpulan
    add_h2(doc, "5.1", "Kesimpulan")
    add_body(doc, "Berdasarkan seluruh tahapan perancangan, implementasi, dan pengujian yang telah dilakukan pada Sistem Informasi Pelayanan Desa (SIPEDES) Rombiyah Barat, dapat ditarik beberapa kesimpulan sebagai berikut:")
    add_body(doc, "1. Telah berhasil dirancang dan dibangun Sistem Informasi Pelayanan Desa Rombiyah Barat berbasis web memanfaatkan framework Laravel 13, Filament v5, dan Livewire yang menyediakan fasilitas pelayanan mandiri (self-service) bagi warga serta panel manajemen terpadu bagi perangkat desa.")
    add_body(doc, "2. Sistem berhasil mengimplementasikan alur pengajuan surat melalui wizard tiga langkah (pemilihan jenis surat, unggah berkas persyaratan, dan konfirmasi) dengan validasi berkas pada sisi server, pembangkitan nomor permohonan otomatis, serta pemantauan status permohonan secara real-time pada halaman riwayat warga dan panel admin.")
    add_body(doc, "3. Sistem berhasil mengintegrasikan asisten virtual berbasis Kecerdasan Buatan dengan arsitektur Retrieval-Augmented Generation (RAG) melalui platform Dify, di mana dokumen pedoman dan SOP pelayanan desa menjadi basis pengetahuan yang menjawab pertanyaan warga secara streaming real-time beserta sumber dokumennya.")
    add_body(doc, "4. Hasil pengujian fungsional menggunakan metode Black-Box Testing otomatis pada 29 skenario kasus uji (TC-01 s/d TC-29) menunjukkan tingkat keberhasilan 100% (PASS) setelah ditemukan dan diperbaiki satu kesalahan (bug) pada penyimpanan berkas permohonan, yang membuktikan bahwa seluruh fitur sistem beroperasi secara optimal, stabil, dan sesuai dengan spesifikasi kebutuhan yang ditetapkan.")

    # 5.2 Saran
    add_h2(doc, "5.2", "Saran")
    add_body(doc, "Untuk pengembangan dan penyempurnaan Sistem Informasi Pelayanan Desa (SIPEDES) Rombiyah Barat di masa yang akan datang, diajukan beberapa saran pengembangan sebagai berikut:")
    add_body(doc, "1. Integrasi Sertifikat Elektronik Resmi (BSrE BSSN): Disarankan untuk mengintegrasikan modul pengesahan dokumen dengan Tanda Tangan Elektronik (TTE) dari Penyelenggara Sertifikat Elektronik (PSrE) resmi pemerintah seperti Balai Sertifikasi Elektronik (BSrE) BSSN agar surat digital yang diterbitkan memiliki kekuatan hukum yang diakui secara nasional sesuai regulasi UU ITE.")
    add_body(doc, "2. Pengembangan Aplikasi Mobile: Perlu dipertimbangkan pengembangan aplikasi mobile berbasis Android/iOS atau Progressive Web App (PWA) untuk lebih memudahkan warga desa dalam mengakses layanan persuratan dan menerima notifikasi status pengajuan langsung di smartphone.")
    add_body(doc, "3. Notifikasi Multi-Kanal: Penambahan mekanisme notifikasi otomatis melalui WhatsApp Gateway dan email pada setiap perubahan status permohonan, sehingga warga mendapatkan pembaruan secara proaktif tanpa harus membuka portal secara manual.")
    add_body(doc, "4. Peningkatan Infrastruktur & Redundansi: Untuk menjamin keberlanjutan layanan, disarankan menyediakan server produksi dengan sumber daya yang memadai, pencadangan data (backup) terjadwal, serta redundansi layanan Dify guna mengantisipasi gangguan pada server utama.")

    print(f"Saving document to {OUT}...")
    doc.save(OUT)
    print("Done generating BAB IV & BAB V docx!")

if __name__ == "__main__":
    import json
    pages = None
    if len(sys.argv) > 1 and os.path.exists(sys.argv[1]):
        with open(sys.argv[1], encoding="utf-8-sig") as f:
            pages = {int(k): v for k, v in json.load(f).items()}
        print("Menggunakan peta halaman gambar:", pages)
    generate_docx(pages)
